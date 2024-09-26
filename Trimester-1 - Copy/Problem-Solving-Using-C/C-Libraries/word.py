import docx
from format import get_data

doc = docx.Document()


for i in get_data():
    p = doc.add_paragraph()
    run = p.add_run(i)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = docx.shared.Pt(16)

    for j in get_data()[i]:
        doc.add_paragraph(j['name'], style='List Number')
        doc.add_paragraph("Syntax: " + j['syntax'], style='List Bullet 2')
        doc.add_paragraph("Description: " + j['desc'], style='List Bullet 2')

doc.save("docx-python Tutorial Demo.docx")
